from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Load the template file

def create_style(doc, font_name, font_size):
    """Create a style with the given font name and size."""
    style_name = f"{font_name}_{font_size}"  # Unique name for each style
    styles = doc.styles
    style = styles.add_style(style_name, 1)  # 1 = Paragraph style
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    return style

def replace(doc, placeholder, value, size="Medium",
            line_spacing=0.9, alignment="left", vertical_alignment=None):
    """
    Replace text placeholders within tables and apply predefined styles for font, size, line spacing, and alignment.

    Args:
        doc: The Document object.
        placeholder: The text placeholder to be replaced.
        value: The new text to insert.
        size: Size of the text ("Big", "Medium", "Small").
        line_spacing: Line spacing (default: 0.9).
        alignment: Horizontal alignment ("left", "center", "right", "justify").
        vertical_alignment: Vertical alignment for cells (None, "top", "center", "bottom").
    """

    # Map string alignment to WD_PARAGRAPH_ALIGNMENT
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    # Define font settings based on size
    font_settings = {
        "Big": ("Lucida Sans Typewriter", 17),
        "Medium": ("Lucida Sans Typewriter", 10),
        "Small": ("Calibri", 9)
    }

    # Get font name and size based on the requested size
    font_name, font_size = font_settings.get(size, font_settings["Medium"])

    # Iterate through tables to find and replace placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    # Iterate through paragraphs to find the last one
                    last_paragraph = cell.paragraphs[-1]  # Get the last paragraph

                    # Clear the last paragraph and add new text
                    last_paragraph.clear()  # Clear the existing content
                    run = last_paragraph.add_run(value)  # Add the new text

                    # Set font properties directly on the run
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

                    # Set paragraph formatting
                    last_paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
                    last_paragraph.paragraph_format.line_spacing = line_spacing

                    # Set vertical alignment for the cell, if provided
                    if vertical_alignment:
                        cell.vertical_alignment = getattr(
                            WD_ALIGN_PARAGRAPH, vertical_alignment.upper(), None
                        )

def place_pictograms(doc, placeholder, image_paths,
                     alignment="center", vertical_alignment="center", width_cm=2):
    """
    Replace an image placeholder with a list of images with alignment options.

    Args:
        doc: The Document object.
        placeholder: The text placeholder to be replaced.
        image_paths: List of image paths to insert.
        alignment: Horizontal alignment ("left", "center", "right").
        vertical_alignment: Vertical alignment for tables ("top", "center", "bottom").
        width_cm: Width of the inserted images in centimeters (default: 2).
    """

    # Map string alignment to WD_ALIGN_PARAGRAPH
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    # Handle images inside table cells if placeholder is found there
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = ""  # Clear the placeholder text
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
                    for image_path in image_paths:
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Cm(width_cm))  # Adjust size

                    # Set vertical alignment, if provided
                    if vertical_alignment:
                        set_vertical_alignment(cell, vertical_alignment)

def set_vertical_alignment(cell, alignment):
    """Set vertical alignment for a table cell."""
    align_map = {"top": "top", "center": "center", "bottom": "bottom"}
    tc_pr = cell._element.get_or_add_tcPr()  # Get or create the <w:tcPr> element
    v_align = OxmlElement("w:vAlign")  # Create the <w:vAlign> element
    v_align.set(qn("w:val"), align_map.get(alignment, "top"))  # Set alignment
    tc_pr.append(v_align)

def fill_form(doc, datos):
    replace(doc, "{nombre}", datos["NOMBRE"], size="Big", alignment="center", vertical_alignment="center")
    replace(doc, "{estado}", datos["ESTADO"], size="Big", alignment="center", vertical_alignment="center")
    replace(doc, "{edicion}", datos["EDICION"], size="Medium", alignment="left", vertical_alignment="center")
    replace(doc, "{revision}", datos["REVISION"], size="Medium", alignment="left", vertical_alignment="center")
    replace(doc, "{ficha}", datos["FICHA"], size="Medium", alignment="left", vertical_alignment="center")
    replace(doc, "{vigente}", datos["VIGENTE"], size="Medium", alignment="left", vertical_alignment="center")
    replace(doc, "{responsable}", datos["RESPONSABLE"], size="Medium", alignment="left", vertical_alignment="center")

    replace(doc, "{riesgos}", datos["RIESGOS"], size="Small", alignment="left", vertical_alignment="top")
    replace(doc, "{manipulacion}", datos["MANIPULACION"], size="Small", alignment="left", vertical_alignment="top")
    replace(doc, "{almacenamiento}", datos["ALMACENAMIENTO"], size="Small", alignment="left", vertical_alignment="top")
    replace(doc, "{eliminacion}", datos["ELIMINACION"], size="Small", alignment="left", vertical_alignment="top")
    replace(doc, "{firstaid}", datos["FIRSTAID"], size="Small", alignment="left", vertical_alignment="top")
    replace(doc, "{derrame}", datos["DERRAME"], size="Small", alignment="left", vertical_alignment="top")
    replace(doc, "{fuego}", datos["FUEGO"], size="Small", alignment="left", vertical_alignment="top")

    place_pictograms(doc, "{sga}", datos["SGA"], width_cm=1.8)